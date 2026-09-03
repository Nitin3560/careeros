import io
from xml.sax.saxutils import escape


def generate_docx(content: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    doc.add_heading(content.get("full_name") or "Resume", level=1)

    if content.get("summary"):
        doc.add_paragraph(content["summary"])

    skills = content.get("skills", [])
    if skills:
        doc.add_heading("Skills", level=2)
        skill_names = ", ".join(skill.get("name", "") for skill in skills if skill.get("name"))
        doc.add_paragraph(skill_names)

    experience = content.get("experience", [])
    if experience:
        doc.add_heading("Experience", level=2)
        for exp in experience:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(exp.get("title", ""))
            run.bold = True
            if exp.get("company"):
                paragraph.add_run(f" - {exp['company']}")
            if exp.get("duration"):
                duration = doc.add_paragraph(exp["duration"])
                duration.runs[0].italic = True
                duration.runs[0].font.size = Pt(10)
            for highlight in exp.get("highlights", []):
                if highlight.strip():
                    doc.add_paragraph(highlight, style="List Bullet")

    education = content.get("education", [])
    if education:
        doc.add_heading("Education", level=2)
        for education_item in education:
            line = f"{education_item.get('degree', '')} - {education_item.get('institution', '')}"
            if education_item.get("year"):
                line += f" ({education_item['year']})"
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _pdf_text(value: object) -> str:
    return escape(str(value or ""))


def generate_pdf(content: dict) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        spaceBefore=12,
        spaceAfter=6,
    )

    story = [Paragraph(_pdf_text(content.get("full_name") or "Resume"), styles["Title"])]

    if content.get("summary"):
        story.append(Spacer(1, 8))
        story.append(Paragraph(_pdf_text(content["summary"]), styles["Normal"]))

    skills = content.get("skills", [])
    if skills:
        story.append(Paragraph("Skills", heading_style))
        skill_names = ", ".join(skill.get("name", "") for skill in skills if skill.get("name"))
        story.append(Paragraph(_pdf_text(skill_names), styles["Normal"]))

    experience = content.get("experience", [])
    if experience:
        story.append(Paragraph("Experience", heading_style))
        for exp in experience:
            title_line = f"<b>{_pdf_text(exp.get('title', ''))}</b>"
            if exp.get("company"):
                title_line += f" - {_pdf_text(exp['company'])}"
            story.append(Paragraph(title_line, styles["Normal"]))
            if exp.get("duration"):
                story.append(Paragraph(f"<i>{_pdf_text(exp['duration'])}</i>", styles["Normal"]))

            bullets = [
                ListItem(Paragraph(_pdf_text(highlight), styles["Normal"]))
                for highlight in exp.get("highlights", [])
                if highlight.strip()
            ]
            if bullets:
                story.append(ListFlowable(bullets, bulletType="bullet"))
            story.append(Spacer(1, 6))

    education = content.get("education", [])
    if education:
        story.append(Paragraph("Education", heading_style))
        for education_item in education:
            line = f"{education_item.get('degree', '')} - {education_item.get('institution', '')}"
            if education_item.get("year"):
                line += f" ({education_item['year']})"
            story.append(Paragraph(_pdf_text(line), styles["Normal"]))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()
